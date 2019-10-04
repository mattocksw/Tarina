
import os
import random
import string
import shutil
import re
import colorsys

import lxml.html
from lxml import etree
from docx import Document, opc, oxml
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH 


from helper_functions import *

#helper class for collecting styles for text
class DocText:
    def __init__(self):
        self.text = ""
        self.font = 'Arial' 
        self.font_size = 13 
        self.alignment = "left" #0 for left, 1 for right, 2 center, 3 justify
        self.bold = False 
        self.italic = False 
        self.underline = False
        self.strikethrough = False
        self.subscript = False
        self.superscript = False
        self.color = None
        self.highlight = None
        self.link = ""
        self.bullet = False
        self.number_list = False
        
    def continue_styles(self, other):
        if(other.bold):
            self.bold = True
        if(other.italic):
            self.italic = True
        if(other.underline):
            self.underline = True
        if(other.strikethrough):
            self.strikethrough = True
        if(other.subscript):
            self.subscript = True
        if(other.superscript):
            self.superscript = True

        self.font = other.font
        self.font_size = other.font_size
        self.color = other.color


    #writes text in given paragraph p in document
    def write(self, p):
        
        #add text
        run = p.add_run(self.text)       

        #add font and font size:
        run.font.name = self.font
        run.font.size = Pt(self.font_size)
        run.font.bold = self.bold
        run.font.italic = self.italic
        run.font.underline = self.underline
        run.font.strike = self.strikethrough
        run.font.subscript = self.subscript
        run.font.superscript = self.superscript
        if(self.color):
            run.font.color.rgb = self.color
        if(self.highlight):
            pass
            #run.font.highlight_color = self.highlight
            #not able to set custom color, only preset values.
        if(self.link):            
            add_hyperlink(p, self.link, self.text)


# add styles from element to DocText t
def add_span_styles(html_element, t):
    #get style
    style = get_style(html_element)

    if(style):
        #get font
        font = get_value_from_string(style, "font-family")
        if(font):
            t.font = font
   
        #get font size
        font_size = get_value_from_string(style, "font-size")
        if(font_size):
            t.font_size = int(font_size[:-2]) #remove Pt
                      

        #get alignment
        alignment = get_value_from_string(style, "text-align")
        if(alignment):
            t.alignment = alignment

        #get text color
        hsl_color = get_value_from_string(style, "color")
        if(hsl_color):
            color = re.findall(r'\d+', hsl_color) #get all numbers https://stackoverflow.com/questions/4289331/how-to-extract-numbers-from-a-string-in-python
            rgb_color = colorsys.hls_to_rgb(float(color[0])/360.0, float(color[2])/100.0, float(color[1])/100.0) #colorsys expects values in 0-1 range
            t.color = RGBColor(int(rgb_color[0]*255.0), int(rgb_color[1]*255.0), int(rgb_color[2]*255.0))

        #get text highlight color (not able to set custom color, ignored)
        #hls_color = get_value_from_string(style, "background-color")
        #if(hls_color):
            #color = re.findall(r'\d+', hls_color) #get all numbers https://stackoverflow.com/questions/4289331/how-to-extract-numbers-from-a-string-in-python
            #rgb_color = colorsys.hls_to_rgb(float(color[0])/100.0, float(color[1])/100.0, float(color[2])/100.0) #colorsys expects values in 0-1 range
            #t.highlight = RGBColor(int(rgb_color[0]*100), int(rgb_color[1]*100), int(rgb_color[2]*100))


# add styles from element to DocText t
def add_tag_styles(html_element, t):
    #get link ()
    #if(html_element.tag == "a"):
    #    t.link = decoration.attrib['href']
    
    if(html_element.tag == "strong"):
        t.bold = True

    elif(html_element.tag == "i"):
        t.italic = True

    elif(html_element.tag == "u"):
        t.underline = True

    elif(html_element.tag == "s"):
        t.strikethrough = True

    elif(html_element.tag == "sub"):
        t.subscript = True

    elif(html_element.tag == "sup"):
        t.superscript = True

# add styles from element to DocText t
def remove_tag_styles(html_element, t):
    if(html_element.tag == "strong"):
        t.bold = False

    elif(html_element.tag == "i"):
        t.italic = False

    elif(html_element.tag == "u"):
        t.underline = False

    elif(html_element.tag == "s"):
        t.strikethrough = False

    elif(html_element.tag == "sub"):
        t.subscript = False

    elif(html_element.tag == "sup"):
        t.superscript = False


#find string between tags from string string_form. return string. First character included in string
def get_next_text(start_i, string_form):
     
    end_i = start_i[0] + 1
    #read until child starts
    while string_form[end_i] != "<":
        end_i += 1
    #get string and update index
    s = string_form[start_i[0]:end_i].replace("\xa0", "")
    start_i[0] = end_i
    #return string
    return s

#advance index over <> Expect first character to be <
def skip_tag(start_i, string_form): 

    #read through the tag itself
    while string_form[start_i[0]] != ">":
        start_i[0] += 1

    #advance to next character
    start_i[0] += 1

#advance index in string until tag skipped
def skip_tags(start_i, string_form):
   
    #skip child tag
    skip_tag(start_i, string_form)
    num_open_tags = 1

    #advance until reaches same level tag
    while num_open_tags > 0:        
        if(string_form[start_i[0]] == "<"):
            if string_form[start_i[0]+1] == "/":
                num_open_tags -= 1
            else:
                num_open_tags += 1
            skip_tag(start_i, string_form)
        else:
            start_i[0] += 1

#Wraps naked text in html tags in span for easier processing
#FIX: currenly expects no spaces after tags in the html which comes directly from ckeditor. Might brake    
def wrap_text(html_element):
    
    elements = list(html_element) 
        
    #wrap text in spans if it has children
    if len(elements): 

        #get string from html and remove what's currently there
        string_form = lxml.html.tostring(html_element).decode('utf-8')
        html_element.text = ""    

        #skip parent element tag
        start_i = [0]
        skip_tag(start_i, string_form)

        insert_index = 0
        children_skipped = 0
        while start_i[0] < len(string_form):
            #if starts with string, read it
            if(string_form[start_i[0]] != "<"):        
                s = get_next_text(start_i, string_form)

                #insert the element 
                e = etree.Element("span")
                e.text = s
                if(insert_index == len(list(html_element))):
                    html_element.append(e)
                else:
                    html_element.insert(insert_index, e)
                insert_index+=1

            else: #skip child
                if(children_skipped >= len(elements)):
                    break #break if we encounter a tag after finishing all childen tags already. Should be main parent tag
                else:
                    skip_tags(start_i, string_form)
                    insert_index+=1
                    children_skipped += 1


#reads the text of element and all its children recursively, and creates a list of DocText objects from them

def get_text(html_element, text_list = [], previousText = DocText()):

    #add spans to text
    if(html_element.text):
        wrap_text(html_element)


    #create new style collector
    t = DocText()
    t.continue_styles(previousText)

    if(html_element.text):
        t.text = html_element.text.replace("\xa0", "")
 
    #get styles from span
    add_span_styles(html_element, t)
          
    #get styles from tag
    add_tag_styles(html_element, t)

    text_list.append(t)

    #recursively add styles
    for element in html_element:
        text_list = get_text(element, text_list,t) 

    return text_list


def get_rowspan(table_column):    

    try:
        if not table_column.attrib:
            return 1
        for attrib, value in table_column.items():
            if(attrib == "rowspan"):
                return int(value)
    except:
        return 1

def get_colspan(table_row):
    try:
        if not table_row.attrib:
            return 1
        for attrib, value in table_row.items():
            if(attrib == "colspan"):
                return int(value)
    except:
        return 1
    
#adds alignment to paragraph defined by styles in text_list
def add_alignment(text_list, p):
    #add alignment
    if(text_list):
        if(text_list[0].alignment == "justify"):
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif(text_list[0].alignment == "right"):
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif(text_list[0].alignment == "center"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def get_tbody(element):
    for tag in element:
        if tag.tag == "tbody":
            return tag
    return element[0]

#return figure if element has such direct child. otherwise return input
def get_figure(element):
    for tag in element:
        if tag.tag == "figure":
            return tag
    return element

def iterate_html(html, document, story_folder):
    for element in html:
        
        #handle figure blocks                    
        if element.tag == "figure":  
            #get width of figure in percentage and adjust in cm (a4 has width of 20.1 cm)
            width = Cm(15.0)
            style = get_style(element)
            if style:
                p_width = get_value_from_string(style, "width")
                width = float(p_width) / 100.0 * width                  
                if(width > Cm(15.0)):
                    width = Cm(15.0)

            for figure in element:
                
                #handle tables
                if figure.tag == "table":

                    tbody = get_tbody(figure)
                        
                    
                    #get dimensions
                    rows = 0
                    cols = 0
                    col_once = False
                    for row in tbody:
                        rows += get_rowspan(row)
                        if col_once == False:
                            col_once = True
                            for col in row:
                                cols += get_colspan(col)
                                
                    #make table
                    table = document.add_table(rows = rows, cols = cols)
                    for idr, row in enumerate(tbody):
                        for idc, col in enumerate(row):                                                        
                            
                            cell = table.cell(idr,idc)
                               
                            col = get_figure(col)

                            if col.tag in ["p", "h2","h3","h4", "td"]:                               

                                text_list = get_text(col, [], DocText())

                                if col.tag in ["p","td"]:
                                    p = cell.add_paragraph()
                                elif col.tag == "h2":
                                    p = cell.add_heading('', level=1)
                                elif col.tag == "h3":
                                    p = cell.add_heading('', level=2)
                                elif col.tag == "h4":
                                    p = cell.add_heading('', level=3)

                                add_alignment(text_list, p)

                                    
                                for text in text_list:
                                    text.write(p)
                            
                            elif col.tag in ["ul","ol"]:
                                for list_element in col:
                                    text_list = get_text(list_element, [], DocText())
                                    if(col.tag == "ul"):
                                        p = cell.add_paragraph(style='List Bullet')
                                    elif(col.tag == "ol"):
                                        p = cell.add_paragraph(style='List Number')           
            
                                    for text in text_list:
                                        text.write(p)

                            elif element.tag == "figure":
                                #handle image
                                for image_in_table in col:
                                    if(image_in_table.tag == "img"):
                                        p = cell.add_paragraph()
                                        run = p.add_run()
                                        add_image(image_in_table, width / cols, run, story_folder)

                            
                
                #handle images
                elif figure.tag == "img":                   
                    add_image(figure, width, document, story_folder)

        #handle text
        elif element.tag in ["p", "h2","h3","h4"]:            
            text_list = get_text(element, [], DocText())
            #add main paragraph
            if element.tag == "p":
                p = document.add_paragraph()
            elif element.tag == "h2":
                p = document.add_heading('', level=1)
            elif element.tag == "h3":
                p = document.add_heading('', level=2)
            elif element.tag == "h4":
                p = document.add_heading('', level=3)

            add_alignment(text_list, p)
            
            for text in text_list:
                text.write(p)

        elif element.tag in ["ul","ol"]:
            for list_element in element:
                text_list = get_text(list_element, [], DocText())
                if(element.tag == "ul"):
                    p = document.add_paragraph(style='List Bullet')
                elif(element.tag == "ol"):
                    p = document.add_paragraph(style='List Number')           
            
                for text in text_list:
                    text.write(p)

#converts story and writes docx document in downloads folder
def create_docx(story_name):
    try:
        #get current path  
        story_folder = get_story_folder(story_name)
        path = default_path + story_folder + '/'

        #get gategory map
        categories = get_file_map(path + categories_file)  

        #create empty document
        document = Document()

        #write title
        document.add_heading(story_name, 0)

        #write all categories
        for category, category_folder in categories.items():
           
            #add subheader
            document.add_heading(category, 1)

            #get chapters        
            category_path = path + category_folder + '/'
            item_map = get_file_map(category_path + item_names_file)
      
            #append all chapters together
            for name, folder in item_map.items():
                if(os.path.isfile(category_path + '/' + folder)):
                    with open(category_path + '/' + folder, encoding='utf-8', mode='r') as file:

                        #write chapter title
                        document.add_heading(name, 2)   

                        #travel through html
                        iterate_html(lxml.html.fromstring(file.read()), document, story_folder)                                  
           
        document.save(default_path + 'story.docx')
        return True

    except:
        print("error making html file")
        return False
                

#return the inline style for element
def get_style(element):
    try:
        for attrib, value in element.items():
            if(attrib == "style"):
                return value
    except:
        return ""


#return value for attrib:value pair in string str
def get_value_from_string(str, attrib):
    tmp = str.split(";")
    for s in tmp:
        if s.startswith(attrib):
            result = s.split(":")
            return result[1].replace("%","")
    return ""



#A function that places a hyperlink within a paragraph object. https://github.com/python-openxml/python-docx/issues/384
#
#param paragraph: The paragraph we are adding the hyperlink to.
#param url: A string containing the required url
#param text: The text displayed for the url
#return: The hyperlink object

def add_hyperlink(paragraph, url, text):

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def add_image(img_element, width, document, story_folder):
    try:
        bofero, delime, image_name = img_element.attrib['src'].rpartition('/')
        for filename in os.listdir(default_path + story_folder):                                     
            if filename.startswith(image_name):   
                print(default_path + story_folder + "/" + filename, width, document)
                document.add_picture(default_path + story_folder + "/" + filename, width=width)
    except:
        print("image could not be added to docx")


#create html that includes all folders and files from the story. The result expexts that images are located in assets folder
def create_html(story_name):
    try:
        #get current path  
        story_folder = get_story_folder(story_name)
        path = default_path + story_folder + '/'

        #get gategory map
        categories = get_file_map(path + categories_file)             

        #overwrite previous file
        #head
        html_string = '''   <!DOCTYPE html>
                            <head>
                            <meta charset="utf-8">
                            <meta http-equiv="X-UA-Compatible" content="IE=edge">
                            <meta name="viewport" content="width=device-width,initial-scale=1.0">

                            <style>
                            body {
                                background: rgb(204,204,204); 
                            }
                            page {
                                background: white;
                                display: block;
                                margin: 0 auto;
                                margin-bottom: 0.5cm;
                                box-shadow: 0 0 0.5cm rgba(0,0,0,0.5);
                                padding: 1.0cm;
                            }
                            page[size="A4"] {  
                                width: 21cm;
                            }
                            @media print {
                            body, page {
                                margin: 0;
                                box-shadow: 0;
                                }
                            }
                            img {
                                max-width: 100%;
                                height: auto;
                            }
                            </style>

                            </head> 
                      '''
        #start body
        html_string += '''  <body>
                            <page size="A4">

                       '''
        
        #write title
        html_string += '<h1>' + story_name + '</h1>'
        with open(download_path + 'story.html', encoding='utf-8', mode='w') as output:
            output.write(html_string)

        #write all categories
        for category, category_folder in categories.items():
           
            html_string = '<h2>' + category + '</h2>'

            #get chapters        
            category_path = path + category_folder + '/'
            item_map = get_file_map(category_path + item_names_file)
      
            #append all chapters together
            for name, folder in item_map.items():
                if(os.path.isfile(category_path + '/' + folder)):
                    with open(category_path + '/' + folder, encoding='utf-8', mode='r') as file:

                        #write chapter title
                        html_string += '<h3>' + name + '</h3>'

                        #write chapter contents
                        html_string += file.read()

            #fix links https://stackoverflow.com/questions/19357506/python-find-html-tags-and-replace-their-attributes
            root = lxml.html.fromstring(html_string)
            for el in root.iter('img'):
                bofero, delime, image_name = el.attrib['src'].rpartition('/')
                for filename in os.listdir(default_path + story_folder):                                     
                    if filename.startswith(image_name):
                        el.attrib['src'] = 'assets/' + filename            
            html_string = lxml.html.tostring(root, pretty_print=True).decode('utf-8')

            #write the category to html
            with open(download_path + 'story.html', encoding='utf-8', mode='a') as output:
                output.write(html_string)
        

        #end body
        html_string = '''  
                            </page>
                            </body>
                       '''      
        with open(download_path + 'story.html', encoding='utf-8', mode='a') as output:
            output.write(html_string)

        return True

    except:
        print("error making html file")
        return False